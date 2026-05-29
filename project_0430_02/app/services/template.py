"""
Template Service - Word模板加载和填充
"""

import os
from io import BytesIO
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.doc_types import DOC_TYPE_LABELS

# 模板目录
TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "..", "templates")

# 文档类型到模板目录的映射
TEMPLATE_MAP = {
    "risk_management": "risk_management",
    "product_spec": "product_spec",
    "instruction": "instruction",
    "sop": "sop",
    "design_development_plan": "design",
    "design_input": "design",
    "design_output": "design",
    "design_review": "design",
    "design_verification": "design",
    "design_validation": "design",
    "design_change": "design",
    "design_history_file": "design"
}


class TemplateService:
    """模板服务"""

    def __init__(self, template_dir: Optional[str] = None):
        self.template_dir = template_dir or TEMPLATE_DIR

    def load_template(self, doc_type: str) -> Document:
        """
        根据文档类型加载模板

        Args:
            doc_type: 文档类型

        Returns:
            Document对象
        """
        template_path = self._get_template_path(doc_type)

        if os.path.exists(template_path):
            return Document(template_path)
        else:
            # 如果模板不存在，创建空白文档
            return Document()

    def _get_template_path(self, doc_type: str) -> str:
        """获取模板文件路径"""
        template_subdir = TEMPLATE_MAP.get(doc_type, "default")
        template_path = os.path.join(
            self.template_dir,
            template_subdir,
            "template.docx"
        )
        return template_path

    def fill_template(
        self,
        doc: Document,
        content: str,
        product_name: str,
        doc_type: str
    ) -> Document:
        """
        用AI生成的内容填充模板

        Args:
            doc: Document对象
            content: AI生成的文档内容（Markdown格式）
            product_name: 产品名称
            doc_type: 文档类型

        Returns:
            填充后的Document对象
        """
        # 解析Markdown内容并写入文档
        self._parse_and_fill(doc, content, product_name, doc_type)
        return doc

    def _parse_and_fill(
        self,
        doc: Document,
        content: str,
        product_name: str,
        doc_type: str
    ):
        """解析Markdown内容并填充到文档"""
        # 确保content是字符串
        if not isinstance(content, str):
            content = str(content)

        lines = content.split("\n")

        # 文档标题
        title = doc.add_heading(str(product_name), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加文档类型副标题
        subtitle = doc.add_paragraph(DOC_TYPE_LABELS.get(doc_type, doc_type))
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # 空行

        # 解析Markdown内容
        current_heading = None
        in_code_block = False
        code_content = []
        table_buffer = []  # 收集连续的Markdown表格行

        for line in lines:
            stripped = line.strip()

            # 跳过标题行（第一个#开头的行是标题）
            if stripped.startswith("# ") and current_heading is None:
                continue

            # 代码块处理
            if stripped.startswith("```"):
                self._flush_table(doc, table_buffer)
                if in_code_block:
                    # 结束代码块
                    p = doc.add_paragraph()
                    p.style = "Quote"
                    p.add_run("\n".join(code_content))
                    code_content = []
                in_code_block = not in_code_block
                continue

            if in_code_block:
                code_content.append(stripped)
                continue

            # 表格行处理 - 收集连续表格行到buffer
            if stripped.startswith("| ") or (stripped.startswith("|") and "|" in stripped[1:]):
                table_buffer.append(stripped)
                continue

            # 非表格行，先flush缓冲的表格
            self._flush_table(doc, table_buffer)

            # 标题处理
            if stripped.startswith("## "):
                current_heading = stripped[3:].strip()
                doc.add_heading(current_heading, level=1)
            elif stripped.startswith("### "):
                sub_heading = stripped[4:].strip()
                doc.add_heading(sub_heading, level=2)
            elif stripped.startswith("**") and stripped.endswith("**"):
                # 粗体文本作为重点段落
                p = doc.add_paragraph()
                run = p.add_run(stripped[2:-2])
                run.bold = True
            elif stripped.startswith("- ") or stripped.startswith("* "):
                # 列表项
                p = doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped == "---":
                # 分隔线
                p = doc.add_paragraph()
                p.add_run("─" * 50)
            elif stripped:
                # 普通段落 - 确保是字符串
                doc.add_paragraph(str(stripped))

        # 处理末尾可能残留的表格
        self._flush_table(doc, table_buffer)

    def _flush_table(self, doc: Document, table_buffer: list):
        """将缓冲的Markdown表格行转换为Word表格"""
        if not table_buffer:
            return

        rows = []
        has_header = False

        for row_text in table_buffer:
            cells = [c.strip() for c in row_text.strip("|").split("|")]
            # 检测分隔行 (如 |---|---|)
            is_separator = all(
                c.replace("-", "").replace(":", "").replace(" ", "") == ""
                for c in cells
            )
            if is_separator:
                has_header = True
                continue
            rows.append(cells)

        if not rows:
            table_buffer.clear()
            return

        # 确定列数（取最宽的行）
        num_cols = max(len(row) for row in rows)

        # 填充不足的列
        for row in rows:
            while len(row) < num_cols:
                row.append("")

        # 创建Word表格
        table = doc.add_table(rows=len(rows), cols=num_cols, style="Table Grid")

        for i, row in enumerate(rows):
            for j in range(num_cols):
                cell = table.cell(i, j)
                cell.text = row[j]
                # 表头加粗
                if has_header and i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        # 表格后加空行
        doc.add_paragraph()
        table_buffer.clear()

    def save_document(self, doc: Document, file_path: str):
        """保存文档到文件"""
        doc.save(file_path)

    def document_to_bytes(self, doc: Document) -> bytes:
        """将文档转换为字节流"""
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        # 确保返回的是字节而不是字符串
        return buffer.getvalue()
